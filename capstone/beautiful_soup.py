from bs4 import BeautifulSoup
import requests
import docx



url = 'https://www.dpacnc.com/events'
r = requests.get(url)
html_doc = r.text
src = r.content
# soup = BeautifulSoup(html_doc)
soup = BeautifulSoup(src)

# Find all the tags in HTML
links = soup.find_all("a")   


links_of_shows = []

for h2_tag in soup.find_all('a'):
        t = h2_tag.attrs['href']
        if t.find('detail') > -1:
            links_of_shows.append(t)

links_of_shows = list(set(links_of_shows))



url = 'https://www.dpacnc.com/events/detail/styx-1'
r = requests.get(url)
html_doc = r.text
src = r.content
# soup = BeautifulSoup(html_doc)
soup = BeautifulSoup(src)


for h2_tag in soup.find_all('a'):
        t = h2_tag.attrs['href']
        # print(t)


# WE NEED TO LOOP TO PRINT THE ATTRIBUTES
# EXPORT THEM FOR REVIEW
for h2_tag in soup.find_all('a'):
        print(h2_tag.attrs)
        # print(t)

# tag.string will give us the content in between the tag!
# https://www.youtube.com/watch?v=87Gx3U0BDlo
# https://stevesie.com/apps/ticketmaster-api

# Project Ideas
# https://nycdatascience.com/blog/student-works/web-scraping/web-scraping-and-analysis-of-concert-ticket-resales/


# print(soup.find_all('b'))
# print(soup.p)


# mydoc = docx.Document()
# mydoc.add_paragraph(find_all('p'))
# mydoc.save(r'C:\Users\IoannisLoukakis\desktop\check_html_p.docx')


# links_of_ticketmaster = []

# for h2_tag in soup.find_all('a'):
#         t = h2_tag.attrs['href']
#         if (t.find('ticketmaster') > -1) & (t.find('event') > -1):
#             links_of_ticketmaster.append(t)

# links_of_ticketmaster = list(set(links_of_ticketmaster))




# url = 'https://www.ticketmaster.com/event/2D00583BF510D807?brand=durham&camefrom=CFC_DPAC_STYX'
# r = requests.get(url)
# html_doc = r.text
# src = r.content
# # soup = BeautifulSoup(html_doc)
# soup = BeautifulSoup(src)


# for h2_tag in soup.find_all('a'):
#         t = h2_tag.attrs['href']
#         print(t)

# print(soup)


# mydoc = docx.Document()
# mydoc.add_paragraph(soup.prettify())
# mydoc.save(r'C:\Users\IoannisLoukakis\desktop\check_html_styx.docx')


# links_of_ticketmaster = []

# for h2_tag in soup.find_all('a'):
#         t = h2_tag.attrs['href']
#         if (t.find('ticketmaster') > -1) & (t.find('event') > -1):
#             links_of_ticketmaster.append(t)

# links_of_ticketmaster = list(set(links_of_ticketmaster))





"""
TODO
Load more events
https://www.dpacnc.com/events/special-offers
Get the event description as a text
"""